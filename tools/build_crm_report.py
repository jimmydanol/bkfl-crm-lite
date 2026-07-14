from __future__ import annotations

import argparse
import tempfile
from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


INK = RGBColor(49, 60, 53)
GREEN = RGBColor(61, 92, 67)
DEEP_GREEN = RGBColor(33, 78, 95)
MUTED = RGBColor(102, 112, 105)
GOLD = RGBColor(154, 107, 32)
LIGHT_GREEN = "EEF5EF"
LIGHT_BLUE = "E8F4FD"
LIGHT_GOLD = "FFF8E8"
LIGHT_GRAY = "F2F4F3"
BORDER = "D9E0DA"
WHITE = "FFFFFF"


def set_run_font(run, *, name="Calibri", size=11, color=INK, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa, indent=120):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_hyperlink(paragraph, text, url, color=DEEP_GREEN, bold=False):
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    r_pr.append(r_fonts)
    color_node = OxmlElement("w:color")
    color_node.set(qn("w:val"), str(color))
    r_pr.append(color_node)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    if bold:
        r_pr.append(OxmlElement("w:b"))
    new_run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_page_number(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, separate, value, end):
        run._r.append(node)


def add_body(doc, text, *, bold_lead=None, italic=False, after=6):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.10
    if bold_lead and text.startswith(bold_lead):
        first = paragraph.add_run(bold_lead)
        set_run_font(first, bold=True)
        rest = paragraph.add_run(text[len(bold_lead):])
        set_run_font(rest, italic=italic)
    else:
        run = paragraph.add_run(text)
        set_run_font(run, italic=italic)
    return paragraph


def add_bullet(doc, text, *, level=0):
    paragraph = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    paragraph.paragraph_format.left_indent = Inches(0.5 + level * 0.25)
    paragraph.paragraph_format.first_line_indent = Inches(-0.25)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.line_spacing = 1.10
    run = paragraph.add_run(text)
    set_run_font(run)
    return paragraph


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(text, style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def add_callout(doc, label, text, fill=LIGHT_GREEN, accent=GREEN):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360], indent=120)
    set_table_borders(table, color=BORDER)
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    label_run = paragraph.add_run(label + "  ")
    set_run_font(label_run, color=accent, bold=True)
    text_run = paragraph.add_run(text)
    set_run_font(text_run, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def crop_image(source, target, bottom):
    with Image.open(source) as image:
        bottom = min(bottom, image.height)
        cropped = image.crop((0, 0, image.width, bottom))
        cropped.save(target, optimize=True)


def add_figure(doc, image_path, caption, width=6.25):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(width))
    drawing = run._r.xpath(".//wp:docPr")
    if drawing:
        drawing[0].set("descr", caption)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(3)
    cap.paragraph_format.space_after = Pt(8)
    cap_run = cap.add_run(caption)
    set_run_font(cap_run, size=9.5, color=MUTED, italic=True)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    heading_tokens = {
        1: (16, GREEN, 16, 8),
        2: (13, GREEN, 12, 6),
        3: (12, DEEP_GREEN, 8, 4),
    }
    for level, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Bullet 2"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)

    header = section.header
    header_p = header.paragraphs[0]
    header_p.paragraph_format.space_after = Pt(0)
    header_p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    left = header_p.add_run("BK FASTLANE  |  CRM PRODUCT REPORT")
    set_run_font(left, size=8.5, color=MUTED, bold=True)
    right = header_p.add_run("\tJULY 14, 2026")
    set_run_font(right, size=8.5, color=MUTED, bold=True)

    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_p.paragraph_format.space_before = Pt(0)
    footer_p.paragraph_format.space_after = Pt(0)
    label = footer_p.add_run("BK FastLane - Jimmy demo")
    set_run_font(label, size=8.5, color=MUTED)


def build_report(output, dashboard, contact):
    doc = Document()
    configure_document(doc)
    doc.core_properties.title = "BK FastLane CRM: Contact to Petition Prep"
    doc.core_properties.subject = "Plain-English product report and production recommendation"
    doc.core_properties.author = "BK FastLane"
    doc.core_properties.keywords = "BK FastLane, CRM, bankruptcy intake, petition prep"

    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_before = Pt(18)
    kicker.paragraph_format.space_after = Pt(2)
    kicker_run = kicker.add_run("CRM PRODUCT REPORT")
    set_run_font(kicker_run, size=10, color=GREEN, bold=True)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(5)
    title_run = title.add_run("BK FastLane CRM:\nContact to Petition Prep")
    set_run_font(title_run, size=26, color=DEEP_GREEN, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    subtitle_run = subtitle.add_run("Plain-English product review for Jimmy Stein and Matt McCune")
    set_run_font(subtitle_run, size=13.5, color=MUTED)

    for label, value in (
        ("Date", "July 14, 2026"),
        ("Status", "Live fake-data demo verified"),
        ("Branch", "Jimmy - Contact to Petition Prep"),
    ):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        label_run = p.add_run(label + ": ")
        set_run_font(label_run, size=10.5, bold=True)
        value_run = p.add_run(value)
        set_run_font(value_run, size=10.5, color=MUTED)
    link_p = doc.add_paragraph()
    link_p.paragraph_format.space_after = Pt(14)
    label_run = link_p.add_run("Live demo: ")
    set_run_font(label_run, size=10.5, bold=True)
    add_hyperlink(link_p, "jimmydanol.github.io/bkfl-crm-lite", "https://jimmydanol.github.io/bkfl-crm-lite/", bold=True)

    add_callout(
        doc,
        "VERDICT",
        "The product direction is now clear enough to explain to a solo bankruptcy attorney in a few minutes. Keep this design and make the workflow secure, reliable, and connected before adding more screens.",
    )

    add_heading(doc, "Executive verdict", 1)
    add_body(doc, "The CRM is now centered on one job: move a potential bankruptcy client from first contact to a complete file that can be used to start preparing the petition.")
    add_body(doc, "The record name changes automatically based on what the firm actually has, not on a staff member's guess.")

    lifecycle = doc.add_table(rows=2, cols=3)
    set_table_geometry(lifecycle, [3120, 3120, 3120], indent=120)
    set_table_borders(lifecycle)
    headers = (("LEAD", LIGHT_BLUE, RGBColor(26, 115, 232)), ("CONTACT", LIGHT_GREEN, GREEN), ("MATTER", "E0F2F1", RGBColor(0, 105, 92)))
    descriptions = ("No Intake answers received", "Some Intake data received; answers or documents still need work", "Data and documents complete; ready for petition prep")
    for idx, (label, fill, color) in enumerate(headers):
        shade_cell(lifecycle.cell(0, idx), fill)
        p = lifecycle.cell(0, idx).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(label)
        set_run_font(run, size=11, color=color, bold=True)
        p2 = lifecycle.cell(1, idx).paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(0)
        run2 = p2.add_run(descriptions[idx])
        set_run_font(run2, size=9.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)

    add_heading(doc, "Recommendation", 2)
    add_body(doc, "Use this version as the product-design baseline. Do not add broad case-management or post-filing features yet. First connect this workflow to secure storage, the live Intake, and a controlled pilot.")

    doc.add_page_break()
    add_heading(doc, "What changed", 1)
    changes = [
        "Matt's four-part structure remains, but the four sections are Dashboard, Leads, Contacts, and Matters.",
        "Settings remains available as a small utility instead of competing with daily work.",
        "The Dashboard explains the workflow at a glance and lists the files that need work next.",
        "Contacts appear ahead of Leads because they are closer to petition prep and contain the most valuable staff work.",
        "Each Contact shows Intake-data progress, document progress, and one plain-language next step.",
        "A Contact becomes a Matter only when required data and every document request are complete.",
        "A Matter returns to Contacts if new activity reopens data or document work.",
        "Each debtor file has four focused tabs: Overview, Intake Data, Documents, and Activity.",
        "The document resolved count now includes only approved or excused requests.",
        "Hover explanations and a fake-data-only banner make the demo easier and safer to understand.",
    ]
    for item in changes:
        add_bullet(doc, item)

    with tempfile.TemporaryDirectory(prefix="bkfl_report_images_") as tmp:
        tmp_path = Path(tmp)
        dash_crop = tmp_path / "dashboard.png"
        contact_crop = tmp_path / "contact.png"
        crop_image(dashboard, dash_crop, 790)
        crop_image(contact, contact_crop, 930)

        add_figure(doc, dash_crop, "Figure 1. The live Dashboard gives the firm one prioritized work list and makes the lifecycle visible.", width=6.25)

        doc.add_page_break()
        add_heading(doc, "The intended daily workflow", 1)
        steps = [
            ("1. Create the Lead", "The firm records the person's name and contact information. The person remains a Lead even if the secure Intake link has been sent, because no answers have been received."),
            ("2. Receive Intake data", "As soon as any Intake answers arrive, the person becomes a Contact. This tells the firm there is now a real file to complete."),
            ("3. Complete the Contact", "The Contact screen separates Intake answers from documents. Staff should always be able to answer: What is missing, and who needs to act next?"),
            ("4. Reach petition prep", "When required Intake sections have answers and all document requests are approved or excused, the Contact automatically becomes a Matter. The next action is to begin petition preparation."),
        ]
        for heading, text in steps:
            add_heading(doc, heading, 2)
            add_body(doc, text)

        add_figure(doc, contact_crop, "Figure 2. A Contact record shows the next step and keeps Intake data and document readiness separate.", width=6.25)

    doc.add_page_break()
    add_heading(doc, "What works well now", 1)
    strengths = [
        ("The language matches a small bankruptcy firm.", "Lead, Contact, and Matter describe real changes in the file, not vague sales stages."),
        ("The Dashboard answers what to do next.", "It shows counts, progress, and one short work list instead of acting as a general control center."),
        ("Data and documents stay separate.", "A complete questionnaire does not mean the documents are complete, and complete documents do not cure missing answers."),
        ("A Matter cannot be created too early.", "The file must earn the Matter label from actual readiness."),
        ("Firm judgment remains in control.", "AI may flag or preliminarily accept an upload, but the firm approves, rejects, or excuses it."),
        ("The design stays narrow.", "The CRM stops at a clean, reviewable package instead of trying to replace mature petition-preparation software."),
    ]
    for label, explanation in strengths:
        add_body(doc, f"{label} {explanation}", bold_lead=label)

    add_heading(doc, "What is not production-ready", 1)
    add_callout(doc, "SAFETY BOUNDARY", "This demo is appropriate for product review and synthetic testing. It is not ready for real client data.", fill=LIGHT_GOLD, accent=GOLD)

    add_heading(doc, "Must be completed before a real-client pilot", 2)
    musts = [
        "Real accounts and permissions for attorneys, paralegals, and administrators.",
        "Encrypted client-data and document storage, backups, retention rules, and tested recovery.",
        "A reliable live Intake connection with retries and visible error handling.",
        "Real file controls, malware scanning, and dependable preview and download behavior.",
        "A permanent activity history showing who changed data, approved documents, or sent requests.",
        "Attorney-approved, versioned rules for chapter-specific required answers and documents.",
        "End-to-end synthetic testing of submission, partial completion, follow-up, Matter promotion, reopening, and export.",
    ]
    for item in musts:
        add_bullet(doc, item)

    add_heading(doc, "Important next improvements", 2)
    for item in (
        "Duplicate detection and safe joint-filer handling.",
        "Reliable email and text reminders with consent and delivery history.",
        "Clear ownership so each Contact has a responsible staff member.",
        "Export into the selected petition-preparation system.",
        "Accessibility, keyboard, mobile, performance, and browser testing.",
        "Simple reporting on time waiting for the client versus time waiting for the firm.",
    ):
        add_bullet(doc, item)

    doc.add_page_break()
    add_heading(doc, "Competitive lessons applied", 1)
    competitors = [
        ("BK Questionnaire", "Guided mobile Intake, exact document requests, organized uploads, and reminders show the value of making missing information obvious and the next request easy."),
        ("NextChapter / MyChapter", "A clean debtor-portal handoff keeps client entry separate from firm review and avoids redundant entry."),
        ("Jubilee", "Client data, documents, and communications belong on one debtor file rather than across disconnected tools."),
        ("Stretto Best Case", "Mature petition-preparation software already owns forms preparation and filing. BK FastLane should export or integrate later, not reproduce it here."),
    ]
    for name, lesson in competitors:
        add_heading(doc, name, 2)
        add_body(doc, lesson)

    add_callout(doc, "BEST POSITIONING", "The fastest, clearest way for a solo bankruptcy firm to turn debtor information and documents into a complete petition-prep package.")

    add_heading(doc, "Recommended production order", 1)
    phases = [
        ("Phase 1 - Make the current workflow real", [
            "Confirm the Lead, Contact, and Matter rules with the attorneys.",
            "Define the authoritative Intake and document data model.",
            "Add authentication, permissions, secure storage, and audit history.",
            "Connect the live Intake to the CRM.",
            "Complete the real document-review and follow-up loop.",
        ]),
        ("Phase 2 - Run a controlled synthetic pilot", [
            "Use varied fake clients with missing, inconsistent, and incomplete information.",
            "Confirm staff can always identify what is missing and who acts next.",
            "Test Matter promotion and reopening back to Contacts.",
            "Fix confusion before allowing real client data.",
        ]),
        ("Phase 3 - Limited real-client pilot", [
            "Select a small number of matters with attorney supervision.",
            "Keep a manual fallback and verify every export.",
            "Track errors, staff time, follow-up counts, and time to readiness.",
            "Expand only after handling is reliable and repeatable.",
        ]),
    ]
    for phase, items in phases:
        add_heading(doc, phase, 2)
        for item in items:
            add_bullet(doc, item)

    add_heading(doc, "Measures that matter", 1)
    for item in (
        "Time from Lead creation to first Intake response.",
        "Time from Contact creation to Matter readiness.",
        "Missing Intake sections per Contact.",
        "Documents waiting on the client versus waiting on the firm.",
        "Follow-ups required per Contact.",
        "Matters reopened because work was incomplete.",
        "Staff time spent per file before petition prep.",
    ):
        add_bullet(doc, item)

    add_heading(doc, "Go / no-go recommendation", 1)
    go_table = doc.add_table(rows=1, cols=2)
    decisions = [
        ("Share this fake-data demo", "YES"),
        ("Use this design as the CRM baseline", "YES"),
        ("Begin a controlled synthetic pilot", "YES"),
        ("Store real client information today", "NO"),
        ("Sell it as production-ready software today", "NO"),
    ]
    for _ in range(len(decisions) - 1):
        go_table.add_row()
    set_table_geometry(go_table, [7000, 2360], indent=120)
    set_table_borders(go_table)
    for idx, (decision, result) in enumerate(decisions):
        left, right = go_table.rows[idx].cells
        if idx % 2 == 0:
            shade_cell(left, LIGHT_GRAY)
            shade_cell(right, LIGHT_GRAY)
        lp = left.paragraphs[0]
        lp.paragraph_format.space_after = Pt(0)
        lr = lp.add_run(decision)
        set_run_font(lr, size=10.5, bold=True)
        rp = right.paragraphs[0]
        rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rp.paragraph_format.space_after = Pt(0)
        rr = rp.add_run(result)
        set_run_font(rr, size=10.5, bold=True, color=GREEN if result == "YES" else RGBColor(155, 28, 28))
    doc.add_paragraph().paragraph_format.space_after = Pt(1)

    add_body(doc, "The next value will come from making this exact workflow secure, reliable, and connected - not from adding more screens.")

    add_heading(doc, "Build and verification record", 1)
    build_rows = [
        ("Matt's unchanged base", "a94e1793dac7d08d8e1632f74c2b05b1a9a4a517"),
        ("Jimmy feature branch", "codex/jimmy-contact-to-petition-ready"),
        ("CRM feature commit", "d59accaa6dab9c26e94d261a0c287cf6295f429e"),
        ("GitHub Pages branch", "Jimmy"),
        ("Live verification", "Lifecycle checks, JSX parsing, four rendered routes, Pages build, HTTP response, and browser inspection"),
        ("Matt's branch", "Not changed"),
    ]
    build_table = doc.add_table(rows=1, cols=2)
    for _ in range(len(build_rows) - 1):
        build_table.add_row()
    set_table_geometry(build_table, [2700, 6660], indent=120)
    set_table_borders(build_table)
    for idx, (label, value) in enumerate(build_rows):
        if idx % 2 == 0:
            shade_cell(build_table.cell(idx, 0), LIGHT_GRAY)
            shade_cell(build_table.cell(idx, 1), LIGHT_GRAY)
        p1 = build_table.cell(idx, 0).paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        r1 = p1.add_run(label)
        set_run_font(r1, size=9.5, bold=True)
        p2 = build_table.cell(idx, 1).paragraphs[0]
        p2.paragraph_format.space_after = Pt(0)
        r2 = p2.add_run(value)
        set_run_font(r2, size=9.5, color=MUTED)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)

    add_heading(doc, "Sources", 1)
    source_rows = [
        ("BK Questionnaire", "https://www.bkquestionnaire.com/"),
        ("BK Questionnaire pricing", "https://www.bkquestionnaire.com/public/pricing"),
        ("NextChapter / MyChapter", "https://app.nextchapterbk.com/sign_in_attempts/new"),
        ("Jubilee", "https://www.jubileepro.com/"),
        ("Stretto Best Case", "https://www.stretto.com/bankruptcy-software-services/best-case/"),
    ]
    for label, url in source_rows:
        add_bullet(doc, f"{label}: {url}")

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--output", type=Path, required=True)
    parser.add_argument("--dashboard", type=Path, required=True)
    parser.add_argument("--contact", type=Path, required=True)
    args = parser.parse_args()
    build_report(args.output, args.dashboard, args.contact)
    print(args.output)


if __name__ == "__main__":
    main()
